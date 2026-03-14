"""
export.py
Generate a .docx file from a completed run.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(run: dict) -> bytes:
    """
    Build a Word document from a run dict and return as bytes.
    Structure:
      - Title + company info
      - Coverage summary
      - For each question: Question → Answer → Citation → Confidence
    """
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading("Questionnaire Answers", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    q_name = run.get("questionnaire_name", "")
    if q_name:
        sub = doc.add_paragraph(f"Source: {q_name}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # ── Coverage Summary ───────────────────────────────────────────────────
    cov = run.get("coverage", {})
    if cov:
        doc.add_heading("Coverage Summary", level=2)
        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = "Table Grid"
        rows_data = [
            ("Total Questions",  str(cov.get("total", 0))),
            ("Answered",         str(cov.get("answered", 0))),
            ("Not Found",        str(cov.get("not_found", 0))),
            ("Coverage %",       f"{round(cov.get('answered', 0) / max(cov.get('total', 1), 1) * 100)}%"),
        ]
        for i, (label, val) in enumerate(rows_data):
            tbl.rows[i].cells[0].text = label
            tbl.rows[i].cells[1].text = val
        doc.add_paragraph()

    # ── Q&A Sections ───────────────────────────────────────────────────────
    doc.add_heading("Questions & Answers", level=2)

    results = run.get("results", [])
    for item in results:
        idx = item.get("index", 0) + 1

        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{idx}: {item['question']}")
        q_run.bold = True
        q_run.font.size = Pt(12)

        # Answer
        a_para = doc.add_paragraph()
        a_para.add_run("Answer: ").bold = True
        a_para.add_run(item.get("answer", "Not found in references."))

        # Citation
        c_para = doc.add_paragraph()
        cite_run = c_para.add_run("Citation: ")
        cite_run.bold = True
        cite_run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        c_para.add_run(item.get("citation", "N/A"))

        # Confidence
        conf = item.get("confidence", 0)
        conf_para = doc.add_paragraph()
        conf_run = conf_para.add_run("Confidence: ")
        conf_run.bold = True
        pct = round(conf * 100)
        conf_para.add_run(f"{pct}%")

        doc.add_paragraph()  # spacer between questions

    # ── Save to bytes ──────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
